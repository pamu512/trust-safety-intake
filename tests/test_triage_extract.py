from pathlib import Path

import pytest

from trust_intake.triage_extract import extract_card
from trust_intake.triage_read import read_document, read_sidecar

FIXTURE = Path("tests/fixtures/triage/repeat-claimant.md")


def test_md_extracts_brands_markets_euro(tmp_path, monkeypatch):
    text = FIXTURE.read_text()
    markets = {"SG": ["SG", "Singapore"], "HK": ["HK", "Hong Kong"]}
    card = extract_card(FIXTURE, text, {}, markets)
    assert card["brands"] == ["foodpanda"]
    assert card["markets"] == ["HK", "SG"]
    assert card["euro_impact"]["value"] == 2_500_000
    assert card["rate"]["value"] == 12
    assert card["journey"] in {"claims-cancel", "unknown"}


def test_all_three_brands():
    markets = {}
    card = extract_card(Path("x.md"), "# X\nall three brands\n", {}, markets)
    assert card["brands"] == ["foodora", "foodpanda", "yemeksepeti"]


def test_sidecar_overrides_brands():
    card = extract_card(Path("x.md"), "# X\nfoodora\n", {"brands": ["foodpanda"]}, {})
    assert card["brands"] == ["foodpanda"]


def test_sidecar_brands_filtered_through_brands():
    card = extract_card(Path("x.md"), "# X\nfoodora\n", {"brands": ["acme", "foodpanda"]}, {})
    assert card["brands"] == ["foodpanda"]


def test_html_extracts_title_and_brand(tmp_path):
    path = tmp_path / "note.html"
    path.write_text("<h1>Title</h1><p>foodora Germany</p>", encoding="utf-8")
    text = read_document(path)
    markets = {"DE": ["DE", "Germany"]}
    card = extract_card(path, text, {}, markets)
    assert card["title"] == "Title"
    assert card["brands"] == ["foodora"]
    assert card["markets"] == ["DE"]


def test_docx_extracts_brand_and_euro(tmp_path):
    from docx import Document

    path = tmp_path / "note.docx"
    doc = Document()
    doc.add_heading("Docx Title", level=1)
    doc.add_paragraph("foodpanda Singapore. Exposure €1M.")
    doc.save(path)
    text = read_document(path)
    card = extract_card(path, text, {}, {"SG": ["SG", "Singapore"]})
    assert card["brands"] == ["foodpanda"]
    assert card["markets"] == ["SG"]
    assert card["euro_impact"]["value"] == 1_000_000


def test_estimate_line_sets_flag():
    card = extract_card(Path("x.md"), "# X\nfoodora\n~ €2.5M estimate\n", {}, {})
    assert card["euro_impact"]["value"] == 2_500_000
    assert card["euro_impact"]["source"] == "extract"
    assert card["euro_impact"]["estimate"] is True


def test_bare_account_not_a_journey():
    card = extract_card(Path("x.md"), "# X\nCreate an account for the claimant.\n", {}, {})
    assert card["journey"] == "unknown"


def test_journey_label_sets_account():
    card = extract_card(Path("x.md"), "# X\nJourney: account\n", {}, {})
    assert card["journey"] == "account"


def test_sidecar_string_brands_and_markets_wrapped():
    card = extract_card(
        Path("x.md"),
        "# X\nfoodora\n",
        {"brands": "foodpanda", "markets": "SG"},
        {},
    )
    assert card["brands"] == ["foodpanda"]
    assert card["markets"] == ["SG"]


def test_sidecar_bad_types_no_traceback():
    card = extract_card(Path("x.md"), "# X\nfoodora Singapore\n", {"brands": 1, "markets": 2}, {"SG": ["SG", "Singapore"]})
    assert card["brands"] == ["foodora"]
    assert card["markets"] == ["SG"]


def test_docx_tables_included(tmp_path):
    from docx import Document

    path = tmp_path / "table.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "foodora Singapore. Exposure €500k."
    doc.save(path)
    text = read_document(path)
    card = extract_card(path, text, {}, {"SG": ["SG", "Singapore"]})
    assert card["brands"] == ["foodora"]
    assert card["markets"] == ["SG"]
    assert card["euro_impact"]["value"] == 500_000


def test_sidecar_path_is_stem_meta_json(tmp_path):
    path = tmp_path / "repeat-claimant.md"
    path.write_text("# X\n", encoding="utf-8")
    (tmp_path / "repeat-claimant.meta.json").write_text('{"brands": ["foodora"]}', encoding="utf-8")
    (tmp_path / "repeat-claimant.md.meta.json").write_text('{"brands": ["yemeksepeti"]}', encoding="utf-8")
    assert read_sidecar(path) == {"brands": ["foodora"]}


def test_unsupported_suffix_raises(tmp_path):
    path = tmp_path / "note.xlsx"
    path.write_bytes(b"not-a-brd")
    with pytest.raises(ValueError, match="unsupported"):
        read_document(path)


def test_read_pdf(tmp_path):
    path = tmp_path / "hello.pdf"
    path.write_bytes(
        b"""%PDF-1.1
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R
   /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 48 >>
stream
BT /F1 12 Tf 72 720 Td (foodora) Tj ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000364 00000 n 
trailer
<< /Size 6 /Root 1 0 R >>
startxref
447
%%EOF
"""
    )
    assert "foodora" in read_document(path)
