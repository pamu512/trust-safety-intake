from __future__ import annotations

import json
from html.parser import HTMLParser
from pathlib import Path


class _HTMLText(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []

    def handle_data(self, data: str) -> None:
        self._parts.append(data)

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag == "br":
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li", "tr", "section"}:
            self._parts.append("\n")

    def text(self) -> str:
        return "".join(self._parts)


def read_document(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8")
    if suffix in {".html", ".htm"}:
        parser = _HTMLText()
        parser.feed(path.read_text(encoding="utf-8"))
        return parser.text()
    if suffix == ".docx":
        from docx import Document

        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    if suffix == ".pdf":
        from pypdf import PdfReader

        return "\n".join((page.extract_text() or "") for page in PdfReader(str(path)).pages)
    raise ValueError("unsupported")


def read_sidecar(path: Path) -> dict:
    side = path.with_name(path.stem + ".meta.json")
    if not side.is_file():
        return {}
    raw = json.loads(side.read_text(encoding="utf-8"))
    if not isinstance(raw, dict):
        raise ValueError("sidecar must be a JSON object")
    return raw
